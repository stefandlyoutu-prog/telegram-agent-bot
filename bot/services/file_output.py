"""Генерация файлов для отправки в Telegram: PDF, DOCX, XLSX, STL, CSV, TXT."""

import io
import json
import math
import re
from typing import Any, Dict, List, Optional, Tuple

FILE_INTENT_PATTERN = re.compile(
    r"stl|\.stl|3d[\s-]?печат|3d[\s-]?модел|модел.{0,20}печат|сетк.{0,12}3d|"
    r"бинарн|вложен|прикреп|"
    r"excel|xlsx|\.xlsx|эксель|эксел|таблиц[ауы]?\s*(excel|xlsx)?|"
    r"график|диаграмм|расч[её]тн.{0,15}таблиц|таблиц.{0,15}расч[её]т|"
    r"word|docx|\.docx|ворд|"
    r"pdf|пдф|\.pdf|"
    r"csv|\.csv|"
    r"txt|\.txt|"
    r"файл[ом]?\s*(stl|excel|word|pdf|docx|xlsx)|"
    r"(сделай|создай|выдай|отправь|пришли|скинь|сгенерируй|нужен).{0,30}"
    r"(stl|excel|xlsx|word|docx|pdf|csv|документ|таблиц|вложен)",
    re.IGNORECASE,
)

FILE_REFUSAL_PATTERN = re.compile(
    r"в этом чате.{0,60}не могу|"
    r"не могу.{0,60}(stl|excel|xlsx|word|docx|pdf|csv|файл|документ|таблиц|вложен|бинарн)|"
    r"не могу.{0,40}(физическ|прикрепить|выдать|создать|сгенерировать|отправить).{0,40}"
    r"(stl|excel|word|docx|pdf|файл|вложен|бинарн)|"
    r"не умею.{0,40}(создавать|генерировать|прикреплять).{0,30}файл|"
    r"реально.{0,30}(сгенерировать|прикрепить|отправить).{0,40}(stl|файл|вложен)",
    re.IGNORECASE,
)

FORMAT_LABELS = {
    "pdf": "PDF",
    "docx": "Word (DOCX)",
    "xlsx": "Excel (XLSX)",
    "stl": "STL (3D-печать)",
    "csv": "CSV",
    "txt": "TXT",
}


def wants_file_output(text: Optional[str]) -> bool:
    if not text:
        return False
    return bool(FILE_INTENT_PATTERN.search(text))


COMPLEX_STL_PATTERN = re.compile(
    r"по\s+фото|из\s+фото|как\s+на\s+фото|точн|вымерен|один\s+в\s+один|"
    r"готов.{0,20}(в\s+принтер|к\s+печати)|загрузить.{0,20}(принтер|слайсер)|"
    r"на\s+каждую\s+детал|отдельн.{0,20}детал|сборк|собираются|"
    r"проект|корпус|механизм|генератор|ксеноморф|фигур",
    re.IGNORECASE,
)

SIMPLE_STL_ALLOWED_PATTERN = re.compile(
    r"\bупрощён|\bупрощен|\bчернов|\bтестов|\bтест\b|"
    r"\bпримитив|\bзаготов|\bпрост(?:ой|ая|ое)\s+(?:stl|модел)|"
    r"\bбез\s+точн|\bупрощ",
    re.IGNORECASE,
)


def should_refuse_placeholder_stl(text: Optional[str], *, from_photo: bool = False) -> bool:
    """Не отправлять фейковые STL-примитивы вместо точной модели."""
    if not text:
        return False
    if SIMPLE_STL_ALLOWED_PATTERN.search(text):
        return False
    if from_photo:
        return True
    return bool(COMPLEX_STL_PATTERN.search(text))


def parse_file_count(text: str, default: int = 1) -> int:
    """Сколько файлов просит пользователь (1–10)."""
    if not text:
        return default
    t = text.lower()
    patterns = [
        r"(\d+)\s*(?:stl|файл|файла|файлов|модел|модели|вариант|варианта|вариантов|шт)",
        r"(?:stl|файл|файлов|модел)[^\d]{0,15}(\d+)",
    ]
    for pat in patterns:
        m = re.search(pat, t)
        if m:
            return max(1, min(10, int(m.group(1))))
    if re.search(
        r"на\s+каждую|отдельн.{0,25}(stl|файл)|по\s+одному\s+(stl|файл)|"
        r"кажд.{0,12}детал",
        t,
    ):
        return min(10, max(default, 8))
    if re.search(r"несколько|пачк|набор|комплект", t):
        return min(5, max(default, 3))
    return default


def wants_3d_model_from_photo(text: Optional[str]) -> bool:
    """Запрос 3D-модели / STL с фото (не карточка Авито)."""
    if not text:
        return False
    t = text.lower()
    if re.search(r"карточк|авито|обложк|макет.{0,10}авито", t):
        return False
    return bool(
        re.search(
            r"3d[\s-]?(модел|модель|печат)|stl|для\s+печат|bambu|бамбу|"
            r"слайсер|загрузить\s+в\s+принтер|фигур.{0,20}печат|"
            r"сделай.{0,20}(модел|stl)|точн.{0,10}модел",
            t,
        )
    )


EXPLICIT_STL_FILE_PATTERN = re.compile(
    r"\bstl\b|\.stl|"
    r"(?:пришли|отправь|сделай|создай|выдай|нужен|сгенерируй|скинь).{0,40}(?:stl|\.stl|файл)|"
    r"файл.{0,20}(?:stl|\.stl)",
    re.IGNORECASE,
)


def explicit_stl_file_requested(text: Optional[str]) -> bool:
    """Явный запрос файла .stl — не «bambu» или «для печати» сами по себе."""
    if not text:
        return False
    if not EXPLICIT_STL_FILE_PATTERN.search(text):
        return False
    if should_refuse_placeholder_stl(text):
        return bool(SIMPLE_STL_ALLOWED_PATTERN.search(text))
    return True


def detect_file_format(text: str) -> Optional[str]:
    if not text:
        return None
    t = text.lower()
    if explicit_stl_file_requested(text):
        return "stl"
    if re.search(r"excel|xlsx|\.xlsx|эксель|эксел", t) or (
        "таблиц" in t and re.search(r"excel|xlsx|файл", t)
    ):
        return "xlsx"
    if re.search(r"\bword\b|docx|\.docx|ворд", t) or (
        "документ" in t and "word" in t
    ):
        return "docx"
    if re.search(r"\bcsv\b|\.csv", t):
        return "csv"
    if re.search(r"\btxt\b|\.txt|текстовый файл", t):
        return "txt"
    if re.search(r"pdf|пдф|\.pdf", t):
        return "pdf"
    return None


def resolve_output_file_format(text: str) -> Optional[str]:
    """Определить формат файла для авто-отправки (явные запросы, не «bambu»)."""
    fmt = detect_file_format(text)
    if fmt:
        return fmt
    if not text or not wants_file_output(text):
        return None
    t = text.lower()
    # Не выводим STL из косвенных «3d/печат/bambu» — для этого есть 3D-роутер (Meshy/OpenSCAD/3MF).
    if re.search(r"excel|xlsx|таблиц|эксель|график|диаграмм|расч[её]т", t):
        return "xlsx"
    if re.search(r"word|docx|ворд", t):
        return "docx"
    if re.search(r"csv", t):
        return "csv"
    if re.search(r"pdf|пдф", t):
        return "pdf"
    if re.search(r"txt|текстовый", t):
        return "txt"
    return None


def infer_format_from_refusal(text: str) -> Optional[str]:
    t = (text or "").lower()
    if re.search(r"stl|3d|печат|бинарн|вложен", t):
        return "stl"
    if re.search(r"excel|xlsx|таблиц", t):
        return "xlsx"
    if re.search(r"word|docx", t):
        return "docx"
    if re.search(r"pdf", t):
        return "pdf"
    return None


def looks_like_file_refusal(text: str) -> bool:
    return bool(FILE_REFUSAL_PATTERN.search(text))


def default_filename(fmt: str, title: str = "document") -> str:
    safe = re.sub(r"[^\w\-а-яА-ЯёЁ ]", "", title)[:40].strip() or "document"
    safe = safe.replace(" ", "-")
    ext_map = {"docx": "docx", "xlsx": "xlsx", "stl": "stl", "pdf": "pdf", "csv": "csv", "txt": "txt"}
    ext = ext_map.get(fmt, fmt)
    return f"{safe}.{ext}"


def build_txt(content: str) -> bytes:
    return content.encode("utf-8")


def build_csv(headers: List[str], rows: List[List[str]]) -> bytes:
    import csv

    buf = io.StringIO()
    writer = csv.writer(buf, lineterminator="\n")
    if headers:
        writer.writerow(headers)
    for row in rows:
        writer.writerow(row)
    return buf.getvalue().encode("utf-8-sig")


def build_docx(title: str, sections: List[Tuple[str, str]]) -> bytes:
    from docx import Document

    doc = Document()
    doc.add_heading(title, level=0)
    for heading, body in sections:
        if heading and heading != title:
            doc.add_heading(heading, level=1)
        for para in body.split("\n"):
            p = para.strip()
            if p:
                doc.add_paragraph(p)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_xlsx(sheets: List[Dict[str, Any]]) -> bytes:
    from openpyxl import Workbook

    wb = Workbook()
    wb.remove(wb.active)
    for idx, sheet in enumerate(sheets):
        name = str(sheet.get("name") or f"Лист{idx + 1}")[:31]
        ws = wb.create_sheet(title=name)
        headers = sheet.get("headers") or []
        rows = sheet.get("rows") or []
        if headers:
            ws.append([str(h) for h in headers])
        for row in rows:
            ws.append([str(c) for c in row])
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def _facet_normal(v0, v1, v2) -> Tuple[float, float, float]:
    ux, uy, uz = v1[0] - v0[0], v1[1] - v0[1], v1[2] - v0[2]
    vx, vy, vz = v2[0] - v0[0], v2[1] - v0[1], v2[2] - v0[2]
    nx = uy * vz - uz * vy
    ny = uz * vx - ux * vz
    nz = ux * vy - uy * vx
    length = math.sqrt(nx * nx + ny * ny + nz * nz) or 1.0
    return nx / length, ny / length, nz / length


def _write_facet(lines: List[str], v0, v1, v2) -> None:
    nx, ny, nz = _facet_normal(v0, v1, v2)
    lines.append(f"  facet normal {nx:.6f} {ny:.6f} {nz:.6f}")
    lines.append("    outer loop")
    for v in (v0, v1, v2):
        lines.append(f"      vertex {v[0]:.6f} {v[1]:.6f} {v[2]:.6f}")
    lines.append("    endloop")
    lines.append("  endfacet")


def build_stl_box(width: float, depth: float, height: float) -> bytes:
    w, d, h = width / 2, depth / 2, height / 2
    verts = [
        (-w, -d, -h),
        (w, -d, -h),
        (w, d, -h),
        (-w, d, -h),
        (-w, -d, h),
        (w, -d, h),
        (w, d, h),
        (-w, d, h),
    ]
    faces = [
        (0, 1, 2),
        (0, 2, 3),
        (4, 6, 5),
        (4, 7, 6),
        (0, 4, 5),
        (0, 5, 1),
        (2, 6, 7),
        (2, 7, 3),
        (0, 3, 7),
        (0, 7, 4),
        (1, 5, 6),
        (1, 6, 2),
    ]
    lines = ["solid model"]
    for f in faces:
        _write_facet(lines, verts[f[0]], verts[f[1]], verts[f[2]])
    lines.append("endsolid model")
    return "\n".join(lines).encode("ascii")


def build_stl_cylinder(radius: float, height: float, segments: int = 32) -> bytes:
    segments = max(8, min(64, int(segments)))
    lines = ["solid cylinder"]
    z0, z1 = -height / 2, height / 2
    top_center = (0.0, 0.0, z1)
    bottom_center = (0.0, 0.0, z0)
    ring_top = []
    ring_bottom = []
    for i in range(segments):
        ang = 2 * math.pi * i / segments
        x, y = radius * math.cos(ang), radius * math.sin(ang)
        ring_top.append((x, y, z1))
        ring_bottom.append((x, y, z0))
    for i in range(segments):
        nxt = (i + 1) % segments
        _write_facet(lines, top_center, ring_top[i], ring_top[nxt])
        _write_facet(lines, bottom_center, ring_bottom[nxt], ring_bottom[i])
        _write_facet(
            lines, ring_bottom[i], ring_top[i], ring_top[nxt]
        )
        _write_facet(
            lines, ring_bottom[i], ring_top[nxt], ring_bottom[nxt]
        )
    lines.append("endsolid cylinder")
    return "\n".join(lines).encode("ascii")


def build_stl_sphere(radius: float, segments: int = 16) -> bytes:
    segments = max(6, min(32, int(segments)))
    lines = ["solid sphere"]
    rings: List[List[Tuple[float, float, float]]] = []
    for lat in range(segments + 1):
        phi = math.pi * lat / segments
        ring = []
        for lon in range(segments):
            theta = 2 * math.pi * lon / segments
            ring.append(
                (
                    radius * math.sin(phi) * math.cos(theta),
                    radius * math.sin(phi) * math.sin(theta),
                    radius * math.cos(phi),
                )
            )
        rings.append(ring)
    for lat in range(segments):
        for lon in range(segments):
            nxt = (lon + 1) % segments
            v00, v01 = rings[lat][lon], rings[lat][nxt]
            v10, v11 = rings[lat + 1][lon], rings[lat + 1][nxt]
            _write_facet(lines, v00, v10, v11)
            _write_facet(lines, v00, v11, v01)
    lines.append("endsolid sphere")
    return "\n".join(lines).encode("ascii")


def extract_ascii_stl_from_text(text: str) -> Optional[bytes]:
    block = re.search(r"```(?:stl)?\s*([\s\S]*?)```", text, re.I)
    raw = block.group(1) if block else text
    if "facet normal" in raw.lower() and "vertex" in raw.lower():
        body = raw.strip()
        if not body.lower().startswith("solid"):
            body = "solid model\n" + body
        if "endsolid" not in body.lower():
            body += "\nendsolid model"
        return body.encode("ascii", errors="ignore")
    return None


def _measurement_part_to_spec(part: Dict[str, Any], index: int = 1) -> Dict[str, Any]:
    """Конвертация части из vision-замеров в spec для STL."""
    hint = str(part.get("shape_hint") or part.get("shape") or "").lower()
    h = float(part.get("height_mm") or part.get("height") or 30)
    w = float(part.get("width_mm") or part.get("width") or 0)
    d = float(part.get("depth_mm") or part.get("depth") or 0)
    r = float(part.get("radius_mm") or part.get("radius") or 0)

    if "cyl" in hint or (r > 0 and w <= 0):
        return {
            "shape": "cylinder",
            "radius_mm": r or max(w, d) / 2 or 15 + index * 2,
            "height_mm": h,
            "segments": 32 + (index % 3) * 8,
        }
    if "spher" in hint or "ball" in hint:
        return {
            "shape": "sphere",
            "radius_mm": r or h / 2 or 18 + index,
            "segments": 20 + (index % 4) * 4,
        }
    if "box" in hint or "rect" in hint or w > 0 or d > 0:
        return {
            "shape": "box",
            "width_mm": w or 25 + index * 4,
            "depth_mm": d or 20 + index * 3,
            "height_mm": h,
        }
    shapes = ("cylinder", "box", "sphere")
    shape = shapes[(index - 1) % 3]
    if shape == "cylinder":
        return {"shape": "cylinder", "radius_mm": 12 + index * 3, "height_mm": h, "segments": 36}
    if shape == "sphere":
        return {"shape": "sphere", "radius_mm": 10 + index * 2.5, "segments": 24}
    return {
        "shape": "box",
        "width_mm": 20 + index * 5,
        "depth_mm": 18 + index * 4,
        "height_mm": max(6, h * (0.7 + (index % 3) * 0.15)),
    }


def _vary_spec_for_index(spec: Dict[str, Any], index: int) -> Dict[str, Any]:
    """Уникальные размеры/форма для i-й детали (index с 1)."""
    i = max(1, index)
    s = dict(spec or {})
    shape = str(s.get("shape") or "").lower()
    if not shape or shape == "box" and i % 3 == 1:
        shape = ("cylinder", "box", "sphere")[(i - 1) % 3]
    elif shape == "box" and i % 3 == 2:
        shape = "sphere"
    elif shape == "cylinder" and i % 3 == 0:
        shape = "box"

    scale = 0.72 + (i % 6) * 0.09
    s["shape"] = shape
    if shape == "cylinder":
        s["radius_mm"] = max(
            4.0, round(float(s.get("radius_mm") or 18) * scale + i * 1.5, 1)
        )
        s["height_mm"] = max(
            4.0, round(float(s.get("height_mm") or 35) * (0.85 + (i % 4) * 0.12), 1)
        )
        s["segments"] = int(s.get("segments") or 32) + (i % 5) * 4
        s.pop("width_mm", None)
        s.pop("depth_mm", None)
    elif shape == "sphere":
        s["radius_mm"] = max(
            4.0,
            round(
                float(s.get("radius_mm") or s.get("height_mm") or 20) * scale
                + i,
                1,
            ),
        )
        s["segments"] = int(s.get("segments") or 20) + (i % 4) * 4
    else:
        s["width_mm"] = max(4.0, round(float(s.get("width_mm") or 28) * scale + i * 2, 1))
        s["depth_mm"] = max(
            4.0, round(float(s.get("depth_mm") or 22) * (0.9 + (i % 3) * 0.15), 1)
        )
        s["height_mm"] = max(
            3.0, round(float(s.get("height_mm") or 16) * (0.8 + (i % 5) * 0.1), 1)
        )
    return s


def _stl_items_from_measurements(
    photo_measurements: str, count: int
) -> List[Tuple[bytes, str, str]]:
    data = parse_json_block(photo_measurements)
    if not isinstance(data, dict):
        return []
    parts = data.get("parts")
    if not isinstance(parts, list) or not parts:
        return []

    items: List[Tuple[bytes, str, str]] = []
    seen: set = set()
    for idx, part in enumerate(parts[:count], start=1):
        if not isinstance(part, dict):
            continue
        spec = _vary_spec_for_index(_measurement_part_to_spec(part, idx), idx)
        mesh = build_stl_from_spec(spec)
        attempt = 0
        while hash(mesh) in seen and attempt < 6:
            spec = _vary_spec_for_index(spec, idx + attempt + 2)
            mesh = build_stl_from_spec(spec)
            attempt += 1
        seen.add(hash(mesh))
        name = str(part.get("name") or f"part-{idx:02d}").strip()
        name = re.sub(r"[^\w\-]+", "-", name)[:40] or f"part-{idx:02d}"
        desc = str(part.get("notes") or part.get("description") or "").strip()
        shape_ru = {"cylinder": "цилиндр", "box": "блок", "sphere": "сфера"}.get(
            spec.get("shape", ""), spec.get("shape", "")
        )
        dims = ", ".join(
            f"{k}={v}"
            for k, v in spec.items()
            if k.endswith("_mm") and isinstance(v, (int, float))
        )
        if not desc:
            desc = f"{shape_ru} ({dims})" if dims else shape_ru
        items.append((mesh, name, desc))
    return items


def build_stl_from_spec(spec: Dict[str, Any]) -> bytes:
    shape = str(spec.get("shape") or "box").lower()
    if shape == "ascii_stl" and spec.get("content"):
        return str(spec["content"]).encode("ascii", errors="ignore")
    if shape == "cylinder":
        return build_stl_cylinder(
            float(spec.get("radius_mm") or spec.get("radius") or 20),
            float(spec.get("height_mm") or spec.get("height") or 40),
            int(spec.get("segments") or 32),
        )
    if shape == "sphere":
        return build_stl_sphere(
            float(spec.get("radius_mm") or spec.get("radius") or 25),
            int(spec.get("segments") or 16),
        )
    return build_stl_box(
        float(spec.get("width_mm") or spec.get("width") or 40),
        float(spec.get("depth_mm") or spec.get("depth") or 40),
        float(spec.get("height_mm") or spec.get("height") or 40),
    )


def parse_json_block(text: str) -> Optional[Dict[str, Any]]:
    block = re.search(r"```(?:json)?\s*([\s\S]*?)```", text)
    raw = block.group(1).strip() if block else text.strip()
    start, end = raw.find("{"), raw.rfind("}") + 1
    if start < 0 or end <= start:
        return None
    try:
        return json.loads(raw[start:end])
    except json.JSONDecodeError:
        return None


def parse_xlsx_payload(text: str) -> List[Dict[str, Any]]:
    data = parse_json_block(text)
    if not data:
        return [{"name": "Данные", "headers": ["Поле", "Значение"], "rows": [["—", "—"]]}]
    sheets = data.get("sheets")
    if isinstance(sheets, list) and sheets:
        return sheets
    if "headers" in data or "rows" in data:
        return [data]
    return [{"name": "Данные", "headers": ["Поле", "Значение"], "rows": [["—", "—"]]}]


async def produce_file(
    fmt: str,
    user_request: str,
    context: str,
    text_model: str,
) -> Tuple[bytes, str, str]:
    """Собрать файл: (bytes, filename, caption)."""
    from bot.services import llm
    from bot.services.seo_pdf import build_seo_pdf, parse_sections_from_markdown

    ctx = (context or user_request)[:4000]

    if fmt == "pdf":
        text = await llm.generate_document_markdown(user_request, ctx, text_model)
        sections = parse_sections_from_markdown(text)
        title = sections[0][0] if sections else "Документ"
        data = build_seo_pdf(title, sections)
        return data, default_filename("pdf", title), f"📄 {FORMAT_LABELS['pdf']}"

    if fmt == "docx":
        text = await llm.generate_document_markdown(user_request, ctx, text_model)
        sections = parse_sections_from_markdown(text)
        title = sections[0][0] if sections else "Документ"
        data = build_docx(title, sections)
        return data, default_filename("docx", title), f"📄 {FORMAT_LABELS['docx']}"

    if fmt == "xlsx":
        raw = await llm.generate_xlsx_json(user_request, ctx, text_model)
        sheets = parse_xlsx_payload(raw)
        title = str(sheets[0].get("name") or "Таблица")
        data = build_xlsx(sheets)
        return data, default_filename("xlsx", title), f"📊 {FORMAT_LABELS['xlsx']}"

    if fmt == "stl":
        items = await produce_file_items(
            "stl",
            user_request,
            ctx,
            text_model,
            count=parse_file_count(user_request, 1),
        )
        return items[0]

    if fmt == "csv":
        raw = await llm.generate_xlsx_json(user_request, ctx, text_model)
        sheets = parse_xlsx_payload(raw)
        sheet = sheets[0]
        headers = [str(h) for h in (sheet.get("headers") or [])]
        rows = [[str(c) for c in row] for row in (sheet.get("rows") or [])]
        data = build_csv(headers, rows)
        return data, default_filename("csv", str(sheet.get("name") or "data")), "📋 CSV"

    text = await llm.generate_document_markdown(user_request, ctx, text_model)
    data = build_txt(text)
    return data, default_filename("txt", "document"), "📄 Текстовый файл"


def _stl_items_from_spec(raw: str, user_request: str) -> List[Tuple[bytes, str, str]]:
    """Список (bytes, basename) для STL."""
    data = parse_json_block(raw)
    if isinstance(data, dict) and isinstance(data.get("files"), list):
        specs = data["files"]
    elif isinstance(data, list):
        specs = data
    elif isinstance(data, dict):
        specs = [data]
    else:
        specs = []

    items: List[Tuple[bytes, str, str]] = []
    seen: set = set()
    for idx, spec in enumerate(specs, start=1):
        if not isinstance(spec, dict):
            continue
        spec = _vary_spec_for_index(spec, idx)
        mesh = build_stl_from_spec(spec)
        attempt = 0
        while hash(mesh) in seen and attempt < 6:
            spec = _vary_spec_for_index(spec, idx + attempt + 2)
            mesh = build_stl_from_spec(spec)
            attempt += 1
        seen.add(hash(mesh))
        name = str(spec.get("name") or f"part-{idx:02d}").strip()
        name = re.sub(r"[^\w\-]+", "-", name)[:40] or f"part-{idx:02d}"
        desc = str(spec.get("description") or "").strip()
        items.append((mesh, name, desc))

    if items:
        return items

    stl_ascii = extract_ascii_stl_from_text(raw)
    if stl_ascii:
        base = "alien-figurine" if "ксеноморф" in user_request.lower() else "model"
        return [(stl_ascii, base, "")]

    fallback = {
        "shape": "cylinder",
        "radius_mm": 28,
        "height_mm": 55,
        "segments": 36,
    }
    if "ксеноморф" in user_request.lower() or "alien" in user_request.lower():
        return [
            (
                build_stl_from_spec(
                    {"shape": "cylinder", "radius_mm": 32, "height_mm": 18, "segments": 40}
                ),
                "alien-head",
                "Голова (упрощённый цилиндр)",
            ),
            (
                build_stl_from_spec(
                    {"shape": "cylinder", "radius_mm": 22, "height_mm": 42, "segments": 36}
                ),
                "alien-body",
                "Туловище",
            ),
            (
                build_stl_from_spec(
                    {"shape": "box", "width_mm": 50, "depth_mm": 50, "height_mm": 8}
                ),
                "alien-base",
                "Основание под печать",
            ),
        ]
    return [(build_stl_from_spec(fallback), "model", "Упрощённая модель")]


async def produce_file_items(
    fmt: str,
    user_request: str,
    context: str,
    text_model: str,
    *,
    count: int = 1,
    print_profile: Optional[Dict[str, Any]] = None,
    photo_measurements: Optional[str] = None,
    from_photo: bool = False,
) -> List[Tuple[bytes, str, str]]:
    """Несколько файлов: [(bytes, filename, caption), ...]."""
    from bot.services import llm

    ctx = (context or user_request)[:4000]
    count = max(1, min(10, count))

    if fmt == "stl":
        from bot.services.capabilities import stl_quality_disclaimer

        if photo_measurements:
            stl_parts = _stl_items_from_measurements(photo_measurements, count)
            meas_data = parse_json_block(photo_measurements)
            if isinstance(meas_data, dict) and isinstance(meas_data.get("parts"), list):
                count = max(count, min(10, len(meas_data["parts"])))
        else:
            stl_parts = []

        if len(stl_parts) < count:
            raw = await llm.generate_stl_batch_specs(
                user_request,
                ctx,
                text_model,
                count=count,
                print_profile=print_profile,
                photo_measurements=photo_measurements,
            )
            llm_parts = _stl_items_from_spec(raw, user_request)
            existing_hashes = {hash(m) for m, _, _ in stl_parts}
            for mesh, name, desc in llm_parts:
                if len(stl_parts) >= count:
                    break
                if hash(mesh) in existing_hashes:
                    continue
                existing_hashes.add(hash(mesh))
                stl_parts.append((mesh, name, desc))

        if len(stl_parts) < count:
            base = parse_json_block(photo_measurements or "") or {}
            parts_list = base.get("parts") if isinstance(base, dict) else None
            if isinstance(parts_list, list) and parts_list:
                extra = _stl_items_from_measurements(
                    photo_measurements or json.dumps(base, ensure_ascii=False),
                    count,
                )
                for item in extra:
                    if len(stl_parts) >= count:
                        break
                    if hash(item[0]) not in {hash(m) for m, _, _ in stl_parts}:
                        stl_parts.append(item)

        if len(stl_parts) < count:
            seed = {"shape": "cylinder", "radius_mm": 22, "height_mm": 38, "segments": 36}
            seen = {hash(m) for m, _, _ in stl_parts}
            for i in range(len(stl_parts) + 1, count + 1):
                spec = _vary_spec_for_index(seed, i)
                mesh = build_stl_from_spec(spec)
                attempt = 0
                while hash(mesh) in seen and attempt < 6:
                    spec = _vary_spec_for_index(seed, i + attempt + 3)
                    mesh = build_stl_from_spec(spec)
                    attempt += 1
                seen.add(hash(mesh))
                shape = spec.get("shape", "деталь")
                stl_parts.append(
                    (
                        mesh,
                        f"part-{i:02d}",
                        f"Деталь {i}: {shape}",
                    )
                )

        stl_parts = stl_parts[:count]

        note = stl_quality_disclaimer(from_photo=from_photo)
        prof_line = ""
        if print_profile:
            from bot.services.print_profile import format_profile

            prof_line = f"\n🖨 {format_profile(print_profile)}"
        note = note + prof_line
        out: List[Tuple[bytes, str, str]] = []
        for data, base, part_desc in stl_parts:
            fname = default_filename("stl", base)
            cap = note
            if part_desc:
                cap = f"{note}\n\n📌 {part_desc}"
            out.append((data, fname, cap))
        return out

    single = await produce_file(fmt, user_request, context, text_model)
    return [single]
